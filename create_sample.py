"""Generate sample exam .docx file untuk testing."""
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('Ulangan Harian Bahasa Indonesia', 0)
doc.add_paragraph('Mata Pelajaran: Bahasa Indonesia')
doc.add_paragraph('Kelas: VII')

doc.add_paragraph('Bacalah teks berikut untuk menjawab soal nomor 1-2.')
doc.add_paragraph(
    'Lingkungan yang bersih adalah dambaan semua orang. Lingkungan yang bersih membuat hidup '
    'menjadi nyaman dan sehat. Sebaliknya, lingkungan yang kotor akan mengundang berbagai '
    'penyakit. Oleh karena itu, kita harus selalu menjaga kebersihan lingkungan di sekitar kita.'
)

doc.add_paragraph('1. Apa ide pokok paragraf di atas?')
doc.add_paragraph('A. Lingkungan yang kotor')
doc.add_paragraph('B. Lingkungan yang bersih dambaan semua orang')
doc.add_paragraph('C. Penyakit datang dari lingkungan kotor')
doc.add_paragraph('D. Cara menjaga kebersihan')

doc.add_paragraph('2. Berdasarkan teks, lingkungan kotor akan menyebabkan ...')
doc.add_paragraph('A. nyaman')
doc.add_paragraph('B. sehat')
doc.add_paragraph('C. penyakit')
doc.add_paragraph('D. bahagia')

doc.add_paragraph('3. Pilih semua kalimat yang merupakan kalimat tanya (jawaban bisa lebih dari satu):')
doc.add_paragraph('A. Siapa namamu?')
doc.add_paragraph('B. Saya pergi ke sekolah.')
doc.add_paragraph('C. Berapa umurmu?')
doc.add_paragraph('D. Dia membaca buku.')

doc.add_paragraph('4. Pernyataan: "Ibu kota Indonesia adalah Jakarta." (Benar/Salah)')

doc.add_paragraph('5. Pernyataan: "Matahari terbit dari barat." (Benar/Salah)')

doc.add_paragraph('6. Jodohkan istilah berikut dengan artinya:')
doc.add_paragraph('1. Subjek -> A. Yang menerangkan kata kerja')
doc.add_paragraph('2. Predikat -> B. Pelaku dalam kalimat')
doc.add_paragraph('3. Objek -> C. Kata kerja dalam kalimat')
doc.add_paragraph('4. Keterangan -> D. Yang dikenai pekerjaan')

doc.save('D:/claude/exam-organizer/sample_soal.docx')
print('Sample created: sample_soal.docx')
