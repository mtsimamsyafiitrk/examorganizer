"""Test parser dengan dokumen yang mengandung huruf Arab."""
import sys
sys.path.insert(0, 'D:/claude/exam-organizer')

from docx import Document
from docx_parser import parse_docx, blocks_to_text

# Buat sample docx dengan huruf Arab
doc = Document()
doc.add_heading('Ulangan Bahasa Arab', 0)
doc.add_paragraph('Mata Pelajaran: Bahasa Arab')
doc.add_paragraph('Kelas: VII')

doc.add_paragraph('4. Bentuk huruf "Ta" (ت) di awal kata saat disambung adalah...')
doc.add_paragraph('A. تَ')
doc.add_paragraph('B. تـ')
doc.add_paragraph('C. ـتـ')
doc.add_paragraph('D. ـت')

doc.add_paragraph('5. Manakah penyambungan huruf yang benar untuk kata "Kaa-ta-ba" (كَتَبَ)?')
doc.add_paragraph('A. كَ تَ بَ')
doc.add_paragraph('B. كَتَبَ')
doc.add_paragraph('C. كــتــبَ')
doc.add_paragraph('D. ك ت ب')

doc.add_paragraph('6. Penulisan kata "Madrosatun" (مَدْرَسَةٌ) yang benar adalah...')
doc.add_paragraph('A. مَدْرَسَةٌ')
doc.add_paragraph('B. مَدْرَسَة')
doc.add_paragraph('C. مدرسة')
doc.add_paragraph('D. مَدرَسة')

doc.save('D:/claude/exam-organizer/sample_arab.docx')
print('Sample arab created.')

# Parse
blocks = parse_docx('D:/claude/exam-organizer/sample_arab.docx', 'D:/claude/exam-organizer/extracted_images/test_arab')
print(f"\nTotal blocks: {len(blocks)}")
text = blocks_to_text(blocks)
with open('D:/claude/exam-organizer/parsed_arab_output.txt', 'w', encoding='utf-8') as f:
    f.write(text)
print("\nText tersimpan ke parsed_arab_output.txt")

# Check explicitly that Arabic chars are in the output
arabic_chars = [c for c in text if '؀' <= c <= 'ۿ']
print(f"\nJumlah karakter Arab terdeteksi: {len(arabic_chars)}")
print(f"Total panjang teks: {len(text)} karakter")
